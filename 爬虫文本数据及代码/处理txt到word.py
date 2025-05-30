import os
from docx import Document


def save_txt_files_to_word(folder_path, output_word_path):
    """
    将指定文件夹中的所有txt文件的内容保存到一个Word文档中。

    :param folder_path: 包含txt文件的文件夹路径
    :param output_word_path: 输出Word文档的路径
    """
    # 创建一个新的Word文档
    doc = Document()

    # 遍历文件夹中的所有文件
    for filename in os.listdir(folder_path):
        if filename.endswith('.txt'):
            file_path = os.path.join(folder_path, filename)
            # 读取txt文件的内容
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                # 将内容添加到Word文档中，并添加一个段落
                doc.add_paragraph(content)
                # 可选：在文档中添加一个分隔符，以便区分不同的txt文件内容
                doc.add_page_break()  # 或者你可以使用 doc.add_paragraph("\n\n---------\n\n") 来添加简单的分隔

    # 保存Word文档
    doc.save(output_word_path)


# 使用示例
folder_to_scan = './saved_pages1'
output_word_file = './result.docx'
save_txt_files_to_word(folder_to_scan, output_word_file)